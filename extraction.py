import os
import re
import zipfile
import shutil
import fitz  # PyMuPDF
import pandas as pd
from concurrent.futures import ThreadPoolExecutor
from flask import Flask, request, render_template, send_file
from flask_cors import CORS
import multiprocessing
from datetime import datetime

app = Flask(__name__, template_folder='.')
CORS(app)

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
TEMP_FOLDER = "temp_resumes"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# No AI Model Needed - Pure Regex for Ultimate Speed
NLP = None
def get_nlp(): return None


# ==========================================
# PRE-COMPILED REGEX PIPELINE for HIGH SPEED
# ==========================================

FINAL_HEADERS = [
    "File Name",
    "Candidate Name",
    "Candidate Contact Number",
    "Email",
    "Total Years of Experience",
    "Skills",
    "Current Company",
    "Previous Companies",
    "Current Location",
    "Education (UG / PG Degree)",
    "Languages Known",
    "Years Worked in Current Company"
]

EMAIL_REGEX = re.compile(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+')
PHONE_REGEX = re.compile(r'(?:\+?\d[\d\s\-\(\)]{8,20}\d|\(\d{3}\)\s*\d{3}-\d{4})')

EXP_REGEX = re.compile(r'\b(\d{1,2}(?:\.\d{1,2})?)\s*(?:\+?\s*)?(?:years?|yrs?|y)\b.*?(?:of)?\s*(?:exp(?:erience)?)\b', re.IGNORECASE)
DOB_REGEX = re.compile(r'(?:DOB|Date of[ \-]Birth)\s*[:\-]?\s*([\d]{1,2}[/\-\.][\d]{1,2}[/\-\.]([\d]{2,4}))', re.IGNORECASE)

NP_REGEX = re.compile(r'(?:Notice Period|NP)\s*[:\-]?\s*(\d{1,2}\s*(?:days?|months?|weeks?)|Immediate(?:ly)?)', re.IGNORECASE)
CTC_REGEX = re.compile(r'(?:Current CTC|CTC)\s*[:\-]?\s*([\d\.]+\s*(?:LPA|L|lakhs?|Lacs?|k|K))', re.IGNORECASE)
ECTC_REGEX = re.compile(r'(?:Expected CTC|ECTC)\s*[:\-]?\s*([\d\.]+\s*(?:LPA|L|lakhs?|Lacs?|k|K))', re.IGNORECASE)

UG_REGEX = re.compile(r'\b(B\.?E\.?|B\.?Tech|B\.?Sc\.?|B\.?Com\.?(?:\s*\(?H\)?(?:\s*ons)?)?|B\.?B\.?A\.?|B\.?C\.?A\.?|B\.?A\.?|Bachelor(?:s|' + r"'" + r's)?\s*(?:of)?\s*(?:Engineering|Technology|Science|Commerce|Arts|Business|Computer)|M\.?B\.?B\.?S\.?|B\.?D\.?S\.?|B\.?Arch\.?|L\.?L\.?B\.?|C\.?A\.?|C\.?S\.?)\b', re.IGNORECASE)
PG_REGEX = re.compile(r'\b(M\.?E\.?|M\.?Tech|M\.?Sc\.?|M\.?B\.?A\.?|M\.?C\.?A\.?|M\.?A\.?|Master(?:s|' + r"'" + r's)?\s*(?:of)?\s*(?:Engineering|Technology|Science|Commerce|Arts|Business|Computer|Business Administration)|M\.?D\.?|M\.?S\.?|I\.?C\.?W\.?A\.?|C\.?F\.?A\.?)\b', re.IGNORECASE)

# Extended City and Location list for Indian & Global candidates
INDIAN_CITIES = {
    'mumbai': 'Mumbai', 'delhi': 'Delhi', 'bangalore': 'Bangalore', 'bengaluru': 'Bangalore', 'hyderabad': 'Hyderabad', 'ahmedabad': 'Ahmedabad',
    'chennai': 'Chennai', 'kolkata': 'Kolkata', 'surat': 'Surat', 'pune': 'Pune', 'jaipur': 'Jaipur', 'lucknow': 'Lucknow',
    'kanpur': 'Kanpur', 'nagpur': 'Nagpur', 'indore': 'Indore', 'thane': 'Thane', 'bhopal': 'Bhopal', 'visakhapatnam': 'Visakhapatnam',
    'pimpri': 'Pimpri', 'patna': 'Patna', 'vadodara': 'Vadodara', 'ghaziabad': 'Ghaziabad', 'ludhiana': 'Ludhiana', 'agra': 'Agra',
    'nashik': 'Nashik', 'faridabad': 'Faridabad', 'meerut': 'Meerut', 'rajkot': 'Rajkot', 'kalyan': 'Kalyan', 'vasai': 'Vasai',
    'varanasi': 'Varanasi', 'srinagar': 'Srinagar', 'aurangabad': 'Aurangabad', 'dhanbad': 'Dhanbad', 'amritsar': 'Amritsar',
    'navi mumbai': 'Navi Mumbai', 'allahabad': 'Allahabad', 'howrah': 'Howrah', 'gwalior': 'Gwalior', 'jabalpur': 'Jabalpur',
    'coimbatore': 'Coimbatore', 'vijayawada': 'Vijayawada', 'jodhpur': 'Jodhpur', 'madurai': 'Madurai', 'raipur': 'Raipur',
    'kota': 'Kota', 'chandigarh': 'Chandigarh', 'guwahati': 'Guwahati', 'solapur': 'Solapur', 'hubli': 'Hubli', 'bareilly': 'Bareilly',
    'moradabad': 'Moradabad', 'mysore': 'Mysore', 'gurgaon': 'Gurgaon', 'gurugram': 'Gurugram', 'noida': 'Noida', 'kochi': 'Kochi',
    'thiruvananthapuram': 'Thiruvananthapuram', 'bhubaneswar': 'Bhubaneswar', 'salem': 'Salem', 'warangal': 'Warangal', 'guntur': 'Guntur',
    'bhiwandi': 'Bhiwandi', 'saharanpur': 'Saharanpur', 'amravati': 'Amravati', 'bikaner': 'Bikaner', 'nanded': 'Nanded', 'kolhapur': 'Kolhapur',
    'ajmer': 'Ajmer', 'gulbarga': 'Gulbarga', 'jamnagar': 'Jamnagar', 'ujjain': 'Ujjain', 'lonavala': 'Lonavala', 'siliguri': 'Siliguri',
    'erode': 'Erode', 'hosur': 'Hosur', 'tirupur': 'Tirupur', 'vellore': 'Vellore', 'tuticorin': 'Tuticorin', 'nellore': 'Nellore',
    'mangalore': 'Mangalore', 'belgaum': 'Belgaum', 'davangere': 'Davangere', 'shimoga': 'Shimoga', 'tumkur': 'Tumkur', 'raichur': 'Raichur',
    'bidar': 'Bidar', 'hassan': 'Hassan', 'bellary': 'Bellary', 'bijapur': 'Bijapur', 'trivandrum': 'Thiruvananthapuram', 'calicut': 'Kozhikode',
    'kozhikode': 'Kozhikode', 'thrissur': 'Thrissur', 'kollam': 'Kollam', 'palakkad': 'Palakkad', 'kottayam': 'Kottayam', 'kannur': 'Kannur',
    'alappuzha': 'Alappuzha', 'tirupati': 'Tupati', 'kakinada': 'Kakinada', 'kurnool': 'Kurnool', 'kadapa': 'Kadapa', 'anantapur': 'Anantapur',
    'vizianagaram': 'Vizianagaram', 'eluru': 'Eluru', 'ongole': 'Ongole', 'nandyal': 'Nandyal', 'tenali': 'Tenali', 'proddatur': 'Proddatur',
    'adoni': 'Adoni', 'madanapalle': 'Madanapalle', 'machilipatnam': 'Machilipatnam', 'jamshedpur': 'Jamshedpur', 'ranchi': 'Ranchi', 'bokaro': 'Bokaro',
    'deoghar': 'Deoghar', 'hazaribagh': 'Hazaribagh', 'giridih': 'Giridih', 'ramgarh': 'Ramgarh', 'ujsain': 'Ujjain', 'gandhinagar': 'Gandhinagar',
    'bhavnagar': 'Bhavnagar', 'junagadh': 'Junagadh', 'anand': 'Anand', 'navsari': 'Navsari', 'morbi': 'Morbi', 'mehsana': 'Mehsana',
    'bharuch': 'Bharuch', 'vapi': 'Vapi', 'porbandar': 'Porbandar', 'ankleshwar': 'Ankleshwar', 'veraval': 'Veraval', 'bhuj': 'Bhuj',
    'rohtak': 'Rohtak', 'panipat': 'Panipat', 'karnal': 'Karnal', 'sonipat': 'Sonipat', 'hissar': 'Hissar', 'yamunanagar': 'Yamunanagar',
    'panchkula': 'Panchkula', 'ambala': 'Ambala', 'kurukshetra': 'Kurukshetra', 'sirsa': 'Sirsa', 'bahadurgarh': 'Bahadurgarh', 'rewari': 'Rewari',
    'shimla': 'Shimla', 'solan': 'Solan', 'mandi': 'Mandi', 'dharamsala': 'Dharamsala', 'una': 'Una', 'kullu': 'Kullu', 'hamirpur': 'Hamirpur',
    'srinagar': 'Srinagar', 'jammu': 'Jammu', 'anantnag': 'Anantnag', 'baramulla': 'Baramulla', 'udhampur': 'Udhampur', 'kathua': 'Kathua',
    'pondicherry': 'Puducherry', 'puducherry': 'Puducherry', 'banaras': 'Varanasi', 'trichy': 'Tiruchirappalli', 'tiruchirappalli': 'Tiruchirappalli',
    'prayagraj': 'Allahabad', 'dharmapuri': 'Dharmapuri', 'karur': 'Karur', 'namakkal': 'Namakkal', 'thanjavur': 'Thanjavur', 'dindigul': 'Dindigul',
    'pudukkottai': 'Pudukkottai', 'nagapattinam': 'Nagapattinam', 'cuddalore': 'Cuddalore', 'kanchipuram': 'Kanchipuram', 'tiruvannamalai': 'Tiruvannamalai',
    'jalandhar': 'Jalandhar', 'patiala': 'Patiala', 'bathinda': 'Bathinda', 'mohali': 'Mohali', 'pathankot': 'Pathankot', 'hoshiarpur': 'Hoshiarpur',
    'batala': 'Batala', 'moga': 'Moga', 'abohar': 'Abohar', 'malerkotla': 'Malerkotla', 'khanna': 'Khanna', 'phagwara': 'Phagwara',
    'muzaffarpur': 'Muzaffarpur', 'gaya': 'Gaya', 'bhagalpur': 'Bhagalpur', 'bihar sharif': 'Bihar Sharif', 'darbhanga': 'Darbhanga', 'purnia': 'Purnia',
    'ara': 'Ara', 'begusarai': 'Begusarai', 'katihar': 'Katihar', 'munger': 'Munger', 'chapra': 'Chapra', 'danapur': 'Danapur',
    'durgapur': 'Durgapur', 'asansol': 'Asansol', 'siliguri': 'Siliguri', 'kharagpur': 'Kharagpur', 'bardhaman': 'Bardhaman', 'english bazar': 'English Bazar',
    'baharampur': 'Baharampur', 'haldia': 'Haldia', 'shantipur': 'Shantipur', 'ranaghat': 'Ranaghat', 'krishnanagar': 'Krishnanagar',
    'rourkela': 'Rourkela', 'brahmapur': 'Berhampur', 'berhampur': 'Berhampur', 'sambalpur': 'Sambalpur', 'puri': 'Puri', 'balasore': 'Balasore',
    'bhadrak': 'Bhadrak', 'baripada': 'Baripada', 'jharsuguda': 'Jharsuguda', 'dhenkanal': 'Dhenkanal', 'barbil': 'Barbil',
    'bilaspur': 'Bilaspur', 'korba': 'Korba', 'bhilai': 'Bhilai', 'rajnandgaon': 'Rajnandgaon', 'jagdalpur': 'Jagdalpur', 'ambikapur': 'Ambikapur',
    'dhamtari': 'Dhamtari', 'mahansamund': 'Mahasamund', 'panaji': 'Panaji', 'vasco da gama': 'Vasco da Gama', 'margao': 'Margao', 'mapusa': 'Mapusa',
    'ponda': 'Ponda', 'bicholim': 'Bicholim', 'canacona': 'Canacona', 'san francisco': 'San Francisco', 'new york': 'New York', 'london': 'London',
    'dubai': 'Dubai', 'singapore': 'Singapore', 'seattle': 'Seattle', 'austin': 'Austin', 'boston': 'Boston', 'chicago': 'Chicago',
    'palo alto': 'Palo Alto', 'mountain view': 'Mountain View', 'india': 'India', 'usa': 'USA', 'united states': 'USA', 'uk': 'UK',
    'united kingdom': 'UK', 'canada': 'Canada', 'australia': 'Australia',
    'tambaram': 'Chennai', 'avadi': 'Chennai', 'ambattur': 'Chennai', 'pallavaram': 'Chennai', 'tiruvallur': 'Tiruvallur', 'chengalpattu': 'Chengalpattu',
    'sriperumbudur': 'Sriperumbudur', 'gummidipoondi': 'Gummidipoondi', 'kanchipuram': 'Kanchipuram', 'vellore': 'Vellore', 'ranipet': 'Ranipet',
    'tirupattur': 'Tirupattur', 'vaniyambadi': 'Vaniyambadi', 'ambur': 'Ambur', 'gudiyatham': 'Gudiyatham', 'arcot': 'Arcot',
    'thiruvannamalai': 'Tiruvannamalai', 'arani': 'Arani', 'villupuram': 'Villupuram', 'tindivanam': 'Tindivanam', 'kallakurichi': 'Kallakurichi',
    'dharmapuri': 'Dharmapuri', 'hosur': 'Hosur', 'krishnagiri': 'Krishnagiri', 'salem': 'Salem', 'mettur': 'Mettur', 'attur': 'Attur',
    'namakkal': 'Namakkal', 'tiruchengode': 'Tiruchengode', 'rasipuram': 'Rasipuram', 'erode': 'Erode', 'perundurai': 'Perundurai',
    'gobichettipalayam': 'Gobichettipalayam', 'bhavani': 'Bhavani', 'sathyamangalam': 'Sathyamangalam', 'tiruppur': 'Tiruppur', 'tirupur': 'Tiruppur',
    'udumalaipettai': 'Udumalaipettai', 'dharapuram': 'Dharapuram', 'palladam': 'Palladam', 'kangeyam': 'Kangeyam', 'nilgiris': 'Udhagamandalam',
    'ooty': 'Udhagamandalam', 'udhagamandalam': 'Udhagamandalam', 'conoor': 'Coonoor', 'coimbatore': 'Coimbatore', 'pollachi': 'Pollachi',
    'mettupalayam': 'Mettupalayam', 'valparai': 'Valparai', 'sulur': 'Sulur', 'karur': 'Karur', 'kulithalai': 'Kulithalai',
    'pallapatti': 'Pallapatti', 'ariyalur': 'Ariyalur', 'perambalur': 'Perambalur', 'trichy': 'Tiruchirappalli', 'tiruchirappalli': 'Tiruchirappalli',
    'srirangam': 'Tiruchirappalli', 'thuvakudi': 'Tiruchirappalli', 'manapparai': 'Manapparai', 'pudukkottai': 'Pudukkottai', 'karaikudi': 'Karaikudi',
    'aranthangi': 'Aranthangi', 'thanjavur': 'Thanjavur', 'tanjore': 'Thanjavur', 'kumbakonam': 'Kumbakonam', 'pattukkottai': 'Pattukkottai',
    'tiruvarur': 'Tiruvarur', 'mannargudi': 'Mannargudi', 'nagapattinam': 'Nagapattinam', 'mayiladuthurai': 'Mayiladuthurai', 'sirkazhi': 'Sirkazhi',
    'chidambaram': 'Chidambaram', 'neyveli': 'Neyveli', 'cuddalore': 'Cuddalore', 'panruti': 'Panruti', 'virudhachalam': 'Virudhachalam',
    'madurai': 'Madurai', 'tirumangalam': 'Thirumangalam', 'melur': 'Melur', 'usilampatti': 'Usilampatti', 'thethakudi': 'Theni',
    'theni': 'Theni', 'periyakulam': 'Periyakulam', 'bodinayakanur': 'Bodinayakanur', 'cumbum': 'Cumbum', 'dindigul': 'Dindigul',
    'palani': 'Palani', 'kodaikanal': 'Kodaikanal', 'oddanchatram': 'Oddanchatram', 'ramanathapuram': 'Ramanathapuram', 'paramakudi': 'Paramakudi',
    'rameswaram': 'Rameswaram', 'sivagangai': 'Sivagangai', 'karaikudi': 'Karaikudi', 'virudhunagar': 'Virudhunagar', 'rajapalayam': 'Rajapalayam',
    'srivilliputhur': 'Srivilliputhur', 'aruppukkottai': 'Aruppukkottai', 'sattur': 'Sattur', 'tuticorin': 'Thoothukudi', 'thoothukudi': 'Thoothukudi',
    'kovilpatti': 'Kovilpatti', 'tirunelveli': 'Tirunelveli', 'tenkasi': 'Tenkasi', 'sankarankovil': 'Sankarankovil', 'kadayanallur': 'Kadayanallur',
    'ambasamudram': 'Ambasamudram', 'valliyur': 'Valliyur', 'nagercoil': 'Nagercoil', 'kanyakumari': 'Kanyakumari', 'marthandam': 'Marthandam'
}

LANGUAGES = ["English", "Hindi", "Tamil", "Telugu", "Marathi", "Kannada", "Malayalam", "Gujarati", "Bengali", "Punjabi", "Urdu", "French", "German", "Spanish"]
LANG_REGEX = re.compile(r'\b(' + '|'.join(LANGUAGES) + r')\b', re.IGNORECASE)

# Skill Dictionary mapped to Canonical Names (Resolves Synonym Edge Cases)
SKILL_ALIASES = {
    "javascript": "JavaScript", "js": "JavaScript",
    "typescript": "TypeScript", "ts": "TypeScript",
    "react": "React", "react.js": "React", "reactjs": "React",
    "node.js": "Node.js", "nodejs": "Node.js", "node": "Node.js",
    "vue": "Vue.js", "vue.js": "Vue.js", "vuejs": "Vue.js",
    "angular": "Angular", "angularjs": "Angular",
    "next.js": "Next.js", "nextjs": "Next.js",
    "express": "Express", "express.js": "Express", "expressjs": "Express",
    "jquery": "jQuery",
    "c++": "C++", "cpp": "C++", 
    "c#": "C#", "c-sharp": "C#", "csharp": "C#",
    "machine learning": "Machine Learning", "ml": "Machine Learning",
    "deep learning": "Deep Learning", "dl": "Deep Learning",
    "artificial intelligence": "AI", "ai": "AI",
    "nlp": "NLP", "natural language processing": "NLP",
    "computer vision": "Computer Vision", 
    "aws": "AWS", "amazon web services": "AWS",
    "gcp": "GCP", "google cloud": "GCP",
    "azure": "Azure", "microsoft azure": "Azure",
    "sql": "SQL", "mysql": "MySQL", "postgresql": "PostgreSQL", "postgres": "PostgreSQL",
    "nosql": "NoSQL", "mongodb": "MongoDB", "mongo": "MongoDB", "sqlite": "SQLite", "redis": "Redis",
    "docker": "Docker", 
    "kubernetes": "Kubernetes", "k8s": "Kubernetes",
    "ci/cd": "CI/CD", "cicd": "CI/CD", "jenkins": "Jenkins", "terraform": "Terraform",
    "git": "Git", "github": "GitHub", "gitlab": "GitLab",
    "linux": "Linux", "unix": "Unix", "bash": "Bash",
    "python": "Python", 
    "pandas": "Pandas", "numpy": "NumPy", "scikit-learn": "Scikit-Learn", "sklearn": "Scikit-Learn",
    "pytorch": "PyTorch", "tensorflow": "TensorFlow", "keras": "Keras", "matplotlib": "Matplotlib",
    "java": "Java", "spring": "Spring", "spring boot": "Spring Boot",
    "ruby": "Ruby", "ruby on rails": "Ruby on Rails", "rails": "Ruby on Rails",
    "php": "PHP", "laravel": "Laravel",
    "go": "Go", "golang": "Go",
    "rust": "Rust", "swift": "Swift", "kotlin": "Kotlin", "dart": "Dart",
    "html": "HTML", "css": "CSS", "bootstrap": "Bootstrap", "tailwind": "Tailwind", "terraform": "Terraform",
    "sap fico": "SAP FICO", "tally": "Tally", "tally erp": "Tally", "gst": "GST", "taxation": "Taxation", 
    "auditing": "Auditing", "financial analysis": "Financial Analysis", "financial modeling": "Financial Modeling",
    "banking": "Banking", "accounting": "Accounting", "ms excel": "MS Excel", "excel": "MS Excel", "advanced excel": "Advanced Excel",
    "autocad": "AutoCAD", "solidworks": "SolidWorks", "catia": "CATIA", "pro-e": "Pro-E", "proe": "Pro-E", 
    "gd&t": "GD&T", "gdt": "GD&T", "lean manufacturing": "Lean Manufacturing", "six sigma": "Six Sigma",
    "staad pro": "STAAD Pro", "revit": "Revit", "primavera p6": "Primavera P6", "primavera": "Primavera P6", 
    "ms project": "MS Project", "estimation": "Estimation", "quantity surveying": "Quantity Surveying", "is codes": "IS Codes",
    "brand management": "Brand Management", "digital marketing": "Digital Marketing", "market research": "Market Research", 
    "trade marketing": "Trade Marketing", "btl": "BTL Activation", "ms powerpoint": "MS PowerPoint", "google analytics": "Google Analytics",
    "photoshop": "Adobe Photoshop", "illustrator": "Adobe Illustrator", "indesign": "Adobe InDesign", "figma": "Figma", 
    "coreldraw": "CorelDRAW", "premiere pro": "Adobe Premiere Pro", "ui/ux": "UI/UX", "social media creatives": "Social Media Creatives",
    "outbound sales": "Sales", "insurance": "Insurance Products", "lead conversion": "Lead Conversion", "leadsquared": "LeadSquared",
    "negotiation": "Negotiation", "siem": "SIEM", "splunk": "Splunk", "qradar": "QRadar", "threat hunting": "Threat Hunting",
    "incident response": "Incident Response", "vulnerability assessment": "Vulnerability Assessment", "ceh": "CEH Certified",
    "security+": "Security+", "owasp": "OWASP Top 10", "firewall": "Firewall Management", "unity": "Unity 3D",
    "c# scripting": "C# Scripting", "shader graph": "Shader Graph", "photon": "Photon Networking", "agile": "Agile", "scrum": "Scrum",
    "recruitment": "Recruitment & Selection", "payroll": "Payroll Processing", "hrms": "HRMS", "darwinbox": "Darwinbox",
    "onboarding": "Onboarding", "statutory compliance": "Statutory Compliance", "compliance": "Statutory Compliance",
    "employee engagement": "Employee Engagement", "sap hr": "SAP HR"
}
STRICT_C_SKILLS = {"C": "C", "R": "R"}
SORTED_ALIASES = sorted(list(SKILL_ALIASES.keys()), key=len, reverse=True)
SKILL_PATTERN = re.compile(r'\b(' + '|'.join(re.escape(s) for s in SORTED_ALIASES) + r')\b', re.IGNORECASE)
STRICT_PATTERN = re.compile(r'\b(C|R)\b')

# Organizations and Descriptions to ignore during Company extraction
IGNORE_ORGS = {
    'school', 'college', 'university', 'institute', 'pvt ltd', 'private limited', 'inc', 'llc', 'corporation',
    'express', 'react', 'python', 'java', 'aws', 'amazon web services',
    'data scientist', 'software engineer', 'developer', 'account executive', 'lead', 'marketing',
    'boston', 'miami', 'columbia', 'duke', 'university of miami', 'boston university', 'duke university',
    'experince', 'experience', 'education', 'skills', 'profile', 'summary', 'about me', 'contact',
    'project', 'responsibilities', 'achievements', 'internship', 'trainee', 'student', 'analyst', 'manager',
    'senior', 'junior', 'associate', 'lead', 'chief', 'head', 'coordinator', 'specialist', 'assistant', 'officer',
    'scientist', 'engineer', 'developer', 'consultant', 'accountant', 'executive', 'nih', 'grant', 'funding', 'investigator',
    # Added "Action Verbs" to prevent job descriptions from being extracted as companies
    'led', 'built', 'developed', 'managed', 'created', 'designed', 'architected', 'improved', 'implemented', 'performed',
    'researched', 'analyzed', 'conducted', 'monitored', 'maintained', 'supported', 'assisted', 'coordinated', 'mentored',
    'increased', 'reduced', 'saved', 'optimized', 'scaled', 'modeled', 'deployed', 'migrated', 'automated', 'integrated',
    'place', 'date', 'declaration', 'signature', 'statement', 'hereby', 'truthful', 'correct', 'best of my knowledge',
    'b.com', 'bcom', 'm.sc', 'msc', 'b.a', 'ba', 'm.a', 'ma', 'b.e', 'm.e', 'b.tech', 'm.tech', 'mba', 'mca', 'bca',
    'implemented', 'attended', 'handled', 'prepared', 'assisted', 'reconciliation', 'management', 'services', 'coloring', 'styling', 'attending', 'executing',
    'percentage', 'cgpa', 'grade', 'score', 'marks', 'obtained', 'aggregate'
}

# High-Confidence Predefined Organization List (Ensures major names are never missed)
PREDEFINED_ORGS = {
    'Reliance', 'TCS', 'Infosys', 'Wipro', 'HCL', 'Accenture', 'Cognizant', 'Capgemini', 'IBM', 'Google', 'Microsoft', 
    'Amazon', 'Slice', 'Unacademy', 'Byjus', 'Airtel', 'Jio', 'Deloitte', 'PWC', 'EY', 'KPMG', 'HDFC', 'ICICI', 'SBI', 
    'Axis', 'L&T', 'Godrej', 'Tata', 'Tech Mahindra', 'Cognizant', 'Mindtree', 'Oracle', 'Cisco', 'Adobe', 'Swiggy', 
    'Zomato', 'Ola', 'Uber', 'Paytm', 'PhonePe', 'Flipkart', 'Snapdeal', 'Myntra', 'Freshworks', 'Zoho', 'Postman',
    'Tech Mahindra', 'Vistara', 'IndiGo', 'Mahindra', 'Bajaj', 'Adani', 'Vedanta', 'Asian Paints', 'Titan'
}
IGNORE_TITLES = {
    'B.A.', 'B.S.', 'B.B.A.', 'M.A.', 'M.S.', 'MBA', 'B.Tech', 'M.Tech', 'Engineering', 'Economics', 
    'Marketing', 'Design', 'Science', 'Arts', 'Fellow', 'Postdoctoral', 'Computer Science', 'UC Berkeley', 'NYU'
}


def unzip_file(zip_path, extract_to=TEMP_FOLDER):
    if os.path.exists(extract_to):
        shutil.rmtree(extract_to)
    os.makedirs(extract_to, exist_ok=True)
    try:
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(extract_to)
    except zipfile.BadZipFile:
        return None
    return extract_to


NAME_REGEX = re.compile(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2}$')

def sanitize_filename(filename):
    # Remove extension and clean characters for Excel/Windows safety
    name = os.path.splitext(filename)[0]
    # Remove common prefix patterns like 'india_resume_01_' or 'resume_v2_'
    name = re.sub(r'^(?:india_|resume_)?(?:resume_|v\d_|)?\d*(?:_)?', '', name, flags=re.IGNORECASE)
    # Replace common resume junk with spaces
    name = re.sub(r'(_| - | -|-|Resume|CV|Curriculum Vitae|Job|Apply|Freshers|Freelance)', ' ', name, flags=re.IGNORECASE)
    # Remove digits and extra symbols
    name = re.sub(r'[^A-Za-z\s\.]', '', name)
    # Final cleanup: if name is just 'pdf' or 'docx', return empty
    if name.lower() in ['pdf', 'docx', 'doc', 'resume']: return ""
    return ' '.join(name.split()).title()

def get_files(folder):
    files = []
    valid_extensions = (".pdf", ".docx")
    for root, _, filenames in os.walk(folder):
        for f in filenames:
            if f.lower().endswith(valid_extensions) and not f.startswith("~$"):
                files.append(os.path.join(root, f))
    return files


def extract_pdf_data(file_path):
    text = ""
    hidden_emails = []
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
            for link in page.get_links():
                uri = link.get("uri", "")
                if uri and uri.lower().startswith("mailto:"):
                    hidden_emails.append(uri[7:].strip())
        doc.close()
    except Exception:
        pass
    return text, hidden_emails


def extract_docx_data(file_path):
    text = ""
    try:
        import docx
        doc = docx.Document(file_path)
        # 1. Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # 2. Extract from tables (Crucial for modern resumes like Priya Reddy's)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
    except Exception as e:
        print(f"DOCX Extraction Error ({os.path.basename(file_path)}): {e}")
    return text, []


def parse_text(text, hidden_emails, file_name="Not Provided"):
    # Common position titles to ignore or split
    # Common position titles to ignore or split
    ROLE_WORDS = {"Manager", "Scientist", "Developer", "Analyst", "Engineer", "Executive", "Associate", "Lead", "Senior", "Junior", "Coordinator", "Nurse", "RN", "Consultant", "Specialist", "Director", "Product", "Early", "Stage", "GTM", "Gtm"}
    
    # Initialize all default empty values with professional placeholder
    data = {k: "Not Provided" for k in FINAL_HEADERS}
    data["File Name"] = file_name
    
    if not text:
        return data

    # 1. EMail
    emails = EMAIL_REGEX.findall(text)
    if emails:
        data["Email"] = emails[0]
    elif hidden_emails:
        data["Email"] = hidden_emails[0]

    # 2. Phone
    phones = PHONE_REGEX.findall(text)
    for p in phones:
        digits = re.sub(r'\D', '', p)
        if len(digits) >= 10:
            data["Candidate Contact Number"] = p.strip(); break

    # 3. Skills
    found_skills = set()
    for match in SKILL_PATTERN.finditer(text):
        token = match.group(1).lower()
        if token in SKILL_ALIASES:
            found_skills.add(SKILL_ALIASES[token])
    for match in STRICT_PATTERN.finditer(text): 
        token = match.group(1)
        if token in STRICT_C_SKILLS:
            found_skills.add(STRICT_C_SKILLS[token])
    data["Skills"] = ", ".join(sorted(list(found_skills)))

    # 4. Name extraction - try multiple strategies
    name = ""
    
    # Strategy A: Check first 1-2 lines (DOCX resumes often have name as first line)
    first_lines = [l.strip() for l in text.split("\n") if len(l.strip()) > 1][:3]
    for line in first_lines[:2]:
        # Name-like: 2-3 words, all title case, no special chars, short
        words = line.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w):
            # Not a section header or job title
            if not any(kw in line.lower() for kw in ['resume', 'curriculum', 'profile', 'engineer', 'manager', 'developer', 'analyst', 'consultant', 'objective', 'summary', 'skills', 'contact']):
                if NAME_REGEX.match(line) or re.match(r'^[A-Z][a-z]+(?: [A-Z][a-z.]+){1,3}$', line):
                    temp = re.sub(r'^(Mr|Ms|Mrs|Miss|Dr)\.?\s+', '', line, flags=re.IGNORECASE).strip()
                    if 3 < len(temp) < 50:
                        name = temp.title()
                        break

    # Strategy B: NAME_REGEX scan across first 10 lines
    if not name:
        scan_lines = [l.strip() for l in text.split("\n") if len(l.strip()) > 3][:10]
        for line in scan_lines:
            if NAME_REGEX.match(line):
                temp_name = line.strip().title()
                temp_name = re.sub(r'^(Mr|Ms|Mrs|Miss|Dr)\.?\s+', '', temp_name, flags=re.IGNORECASE)
                if not any(t.lower() in temp_name.lower() for t in ["Manager", "Scientist", "Developer", "Analyst", "Engineer", "Objective", "Summary"]):
                    name = temp_name
                    break

    # Strategy C: Filename fallback
    if not name:
        name = sanitize_filename(data.get('File Name', '')) or "Not Provided"

    data["Candidate Name"] = name.strip()

    # 4b. Location (Search entire text)
    text_lower = text.lower()
    for city, _ in INDIAN_CITIES.items():
        if re.search(r'\b' + re.escape(city) + r'\b', text_lower):
            data["Current Location"] = city.title(); break





    # 5. Experience (Direct extraction or total calculation)
    exp_matches = EXP_REGEX.search(text)
    if exp_matches:
        data["Total Years of Experience"] = f"{exp_matches.group(1)} Years"
    else:
        # Step 5 Fallback: Sum up years from date ranges (Interval Merging Algorithm)
        date_ranges_raw = re.findall(r'\b(\d{4})\s*(?:-|\u2013|\u2014|to)\s*(\d{4}|Present|Current|Today|Till Date)\b', text, re.IGNORECASE)
        intervals = []
        curr_year = datetime.now().year
        for start, end in date_ranges_raw:
            s_yr = int(start)
            e_yr = curr_year if re.search(r'Present|Current|Today|Till Date', end, re.IGNORECASE) else int(end)
            if 1970 < s_yr <= curr_year and 1970 < e_yr <= curr_year and s_yr <= e_yr:
                intervals.append([s_yr, e_yr])
        
        if intervals:
            # Merge overlapping intervals
            intervals.sort(key=lambda x: x[0])
            merged = []
            for interval in intervals:
                if not merged or interval[0] > merged[-1][1]:
                    merged.append(interval)
                else:
                    merged[-1][1] = max(merged[-1][1], interval[1])
            
            total_yrs = sum(itv[1] - itv[0] for itv in merged)
            if total_yrs > 0:
                data["Total Years of Experience"] = f"{total_yrs} Years (Est.)"
        elif re.search(r'\b(fresher|fresh graduate|no experience|entry.?level|0 year)\b', text, re.IGNORECASE):
            data["Total Years of Experience"] = "Fresher"


    # 6. Location (Redundant search removed)

    # 7. Known Languages
    found_langs = set()
    for match in LANG_REGEX.finditer(text):
        found_langs.add(match.group(1).title())
    if found_langs:
        data["Languages Known"] = ", ".join(sorted(list(found_langs)))

    # 8. Degrees
    degrees = []
    ug_match = UG_REGEX.search(text)
    if ug_match:
        degrees.append(ug_match.group(1).strip())
    pg_match = PG_REGEX.search(text)
    if pg_match:
        degrees.append(pg_match.group(1).strip())
    if degrees:
        data["Education (UG / PG Degree)"] = " / ".join(degrees)

    # 9. Companies & Years Worked in Current Company
    exp_idx = re.search(r'\b(?:Experience|Employment History|Work Experience)\b', text, re.IGNORECASE)
    if exp_idx:
        post_exp_text = text[exp_idx.end():exp_idx.end() + 4000]

        # Calculate Years Worked in Current Company
        # Match year ranges: "Jan 2018 - 2022", "(2021 to Present)", support –, —, -
        date_pattern = r'\(?((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?[a-z]*[\s\,]*\d{4})\s*(?:-|–|—|to)\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?[a-z]*[\s\,]*\d{4}|Present|Current|Till Date|Today)\)?'
        first_date_range = re.search(date_pattern, post_exp_text, re.IGNORECASE)
        if first_date_range:
            start_str, end_str = first_date_range.groups()
            start_yr_match = re.search(r'(\d{4})', start_str)
            end_yr_match = re.search(r'(\d{4})', end_str)

            if start_yr_match:
                start_yr = int(start_yr_match.group(1))
                if re.search(r'Present|Current|Till|Today', end_str, re.IGNORECASE):
                    end_yr = datetime.now().year
                elif end_yr_match:
                    end_yr = int(end_yr_match.group(1))
                else:
                    end_yr = start_yr
                
                yr_diff = end_yr - start_yr
                if yr_diff == 0:
                    data["Years Worked in Current Company"] = "< 1 Year"
                elif 0 < yr_diff < 40:
                    data["Years Worked in Current Company"] = f"{yr_diff}"

        # ============================================================
        # 9c. DATE-ANCHORED Company Extraction Engine
        # Works with ALL resume templates:
        #   Format 1: "Infosys Ltd, Hyderabad – Engineer (Aug 2021 – Present)"
        #   Format 2: "Infosys Ltd\nAug 2021 – Present"
        #   Format 3: "Infosys Ltd\nSystems Engineer\n2021 – Present"
        #   Format 4: "Engineer at Infosys Ltd (2021 – Present)"
        #   Format 5: "Infosys Ltd | Engineer | 2021–Present"
        # ============================================================

        lines = [l.strip() for l in post_exp_text.split("\n") if len(l.strip()) > 2][:60]
        found_orgs_with_dates = []
        found_orgs_fallback = []

        ACTION_VERBS = {
            'led', 'built', 'developed', 'managed', 'created', 'designed', 'architected',
            'improved', 'implemented', 'performed', 'researched', 'analyzed', 'conducted',
            'monitored', 'maintained', 'supported', 'assisted', 'coordinated', 'mentored',
            'increased', 'reduced', 'saved', 'optimized', 'scaled', 'modeled', 'deployed',
            'migrated', 'automated', 'integrated', 'achieved', 'raised', 'prepared',
            'executed', 'grew', 'supervised', 'implement', 'handling', 'preparation',
            'assisting', 'attending', 'managing', 'published', 'trained', 'conducted',
            'specialized', 'consistently', 'provided', 'reported', 'collaborated',
            'spearheaded', 'oversaw', 'coordinated', 'processed', 'evaluated'
        }

        EDU_JUNK = {
            'bcom', 'msc', 'mba', 'btech', 'mtech', 'bca', 'mca', 'ba', 'ma', 'bsc',
            'msc', 'phd', 'diploma', 'account', 'percentage', 'cgpa', 'grade', 'score',
            'marks', 'aggregate', 'gpa', 'class'
        }

        COMPANY_KW = r'\b(?:Pvt|Ltd|Limited|Corp|Inc|Company|Agency|Solutions|Firm|Bank|'\
                     r'Systems|Research|Financial|Hospital|Medical|Clinic|Center|Health|'\
                     r'Services|Tech|Software|Technologies|Consulting|Consultancy|'\
                     r'Industries|Group|Associates|Enterprises|International|Global|'\
                     r'Manufacturing|Analytics|Digital|Networks|Media|Studios)\b'

        # Stricter Date Pattern (Handles 2023, '23, 23 after month)
        STRICT_DATE_RE = re.compile(r'\b(20\d{2}|199\d|(?<=\s)\'(2[0-5])|(?<=\s)(2[0-5])|Present|Current|Till|Today)\b', re.IGNORECASE)

        # Full date range pattern (Improved for 2-digit years)
        DATE_RANGE_RE = re.compile(
            r'(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s,\']*(\d{2,4})?|(\d{4}))\s*(?:-|\u2013|\u2014|to)\s*'
            r'(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s,\']*(\d{2,4})?|(\d{4}|Present|Current|Till Date|Today))',
            re.IGNORECASE
        )
        DATE_YEAR_ONLY_RE = re.compile(r'\b(\d{4}|Present|Current)\s*[-–]\s*(\d{4}|Present|Current)\b', re.IGNORECASE)

        def clean_company(raw):
            """Clean raw text into a company name."""
            c = re.sub(r'^[\u2022\-\*\d\.\s]+', '', raw).strip()
            c = re.sub(r'\(.*?\)', '', c).strip()  # Remove parentheses
            c = re.sub(r'\s*\b(Present|Current|Till Date|Today)\b.*', '', c, flags=re.IGNORECASE)
            for rw in ROLE_WORDS:
                c = re.sub(r'\b' + re.escape(rw) + r'\b', '', c, flags=re.IGNORECASE)
            c = re.sub(r'199\d|20\d{2}', '', c)  # Remove years
            # Strip trailing Indian city names (e.g. 'Infosys Ltd Hyderabad' -> 'Infosys Ltd')
            city_suffix = r'\b(' + '|'.join(re.escape(city.title()) for city in list(INDIAN_CITIES.keys())[:60]) + r')\s*$'
            c = re.sub(city_suffix, '', c, flags=re.IGNORECASE).strip()
            c = re.sub(r'[^A-Za-z\s\.\&\-\'\(\)]', '', c).strip()
            c = re.sub(r'\s*[\u2013\u2014\-]\s*$', '', c).strip()  # Trailing dashes
            return ' '.join(c.split())

        def score_candidate(text):
            """Score how likely a text fragment is a company name."""
            if not text or len(text) < 2:
                return 0
            score = 0
            txt_lower = text.lower()
            
            # --- POSITIVE SIGNALS ---
            if re.search(COMPANY_KW, text, re.IGNORECASE): score += 50
            if any(po.lower() in txt_lower for po in PREDEFINED_ORGS): score += 60
            if text[0].isupper(): score += 10
            
            words = text.split()
            if words and sum(1 for w in words if w and w[0].isupper()) == len(words):
                score += 15
            
            # Length calibration
            if len(words) <= 4: score += 15
            elif len(words) <= 6: score += 5
            elif len(words) > 8: score -= 30
            
            # --- NEGATIVE SIGNALS (PENALITIES) ---
            if any(re.search(r'\b' + v + r'\b', txt_lower) for v in ACTION_VERBS):
                score -= 100
            
            # Case 1 & 2: Projects / Internships / Academic
            PROJECT_WORDS = {'project', 'intern', 'training', 'student', 'academic', 'thesis', 'curriculum', 'education', 'university', 'college'}
            if any(re.search(r'\b' + w + r'\b', txt_lower) for w in PROJECT_WORDS):
                score -= 120
            
            SKILL_SIGNALS = {'excel', 'tally', 'python', 'java', 'sql', 'aws', 'sap', 'powerpoint', 'photoshop', 'autocad', 'figma', 'r', 'ms'}
            if any(re.search(r'\b' + s + r'\b', txt_lower) for s in SKILL_SIGNALS):
                score -= 80
            
            if any(w in txt_lower for w in IGNORE_ORGS): score -= 60
            
            first_w = re.sub(r'[^a-z]', '', words[0].lower()) if words else ''
            if first_w in EDU_JUNK: score -= 100
            
            if DATE_RANGE_RE.match(text.strip()) or DATE_YEAR_ONLY_RE.match(text.strip()):
                score -= 100
            
            if text[0].islower(): score -= 100
            if len(text) > 50: score -= 40
            
            return score

        # Step 1: Find all date range lines and their positions
        date_line_set = set()
        for i, line in enumerate(lines):
            if DATE_RANGE_RE.search(line) or re.search(r'\b(Present|Current|Today|Till Date)\b', line, re.IGNORECASE):
                date_line_set.add(i)

        # Step 2: For each date range, find the best company candidate in surrounding lines
        already_extracted = set()
        for date_idx in sorted(date_line_set):
            date_line = lines[date_idx]
            is_current = bool(re.search(r'\b(Present|Current|Today|Till Date)\b', date_line, re.IGNORECASE))

            # Case 5: Increased window for Table Layouts
            # Collect candidate lines: current line + 6 lines before + 1 line after
            lookback = max(0, date_idx - 6)
            candidate_lines = lines[lookback:date_idx + 1]

            best_name = None
            best_score = 0

            for cline in candidate_lines:
                if any(term in cline.lower() for term in ['declaration', 'signature', 'hobbies', 'interest']):
                    break

                # Skip pure date lines
                if DATE_RANGE_RE.match(cline.strip()):
                    continue

                # --- Strategy A: Line has a date inline (e.g. "Infosys Ltd – Engineer (Aug 2021)") ---
                if STRICT_DATE_RE.search(cline):
                    # Extract pre-date part
                    pre_date = re.split(r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{4})', cline, flags=re.IGNORECASE)[0]
                    # Also try splitting on dash/en-dash before date
                    pre_dash = re.split(r'\s*(?:[\u2013\u2014]|-{1,2})\s*(?=[A-Z]|\d{4}|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)', cline)[0]
                    # Use the shorter of the two (more likely to be just the company)
                    raw = pre_date if len(pre_date) < len(pre_dash) else pre_dash
                    cleaned = clean_company(raw)
                    # Remove location after comma: "Infosys Ltd, Hyderabad" → "Infosys Ltd"
                    if ',' in cleaned:
                        cleaned = cleaned.split(',')[0].strip()
                    s = score_candidate(cleaned)
                    if s > best_score and cleaned not in already_extracted:
                        best_score = s
                        best_name = cleaned

                # --- Strategy B: "at" format (e.g. "Engineer at Infosys") ---
                at_match = re.search(r'\bat\b(.+?)(?:,|\(|$)', cline, re.IGNORECASE)
                if at_match:
                    raw = at_match.group(1).strip()
                    cleaned = clean_company(raw)
                    s = score_candidate(cleaned)
                    if s > best_score and cleaned not in already_extracted:
                        best_score = s
                        best_name = cleaned

                # --- Strategy C: Pipe-separated line ("Company | Role | Date") ---
                if '|' in cline:
                    parts = [p.strip() for p in cline.split('|')]
                    for part in parts:
                        cleaned = clean_company(part)
                        s = score_candidate(cleaned)
                        if s > best_score and cleaned not in already_extracted:
                            best_score = s
                            best_name = cleaned

                # --- Strategy D: Plain company line (separate line from date) ---
                cleaned = clean_company(cline)
                # Remove comma-separated location
                if ',' in cleaned:
                    cleaned = cleaned.split(',')[0].strip()
                s = score_candidate(cleaned)
                if s > best_score and cleaned not in already_extracted:
                    best_score = s
                    best_name = cleaned

            # Accept the candidate if score is above threshold
            if best_name and best_score >= 20 and len(best_name) >= 2:
                # Canonicalize predefined brands
                for po in PREDEFINED_ORGS:
                    if po.lower() in best_name.lower():
                        best_name = po
                        break

                already_extracted.add(best_name)
                if is_current:
                    found_orgs_with_dates.append(best_name)
                else:
                    found_orgs_fallback.append(best_name)

        # 10. Assign based on confidence
        current_company = "Not Provided"
        previous_list = []

        if found_orgs_with_dates:
            current_company = found_orgs_with_dates[0]
            other_potential = found_orgs_with_dates[1:] + found_orgs_fallback
        elif found_orgs_fallback:
            current_company = found_orgs_fallback[0]
            other_potential = found_orgs_fallback[1:]
        else:
            other_potential = []

        # Freelance/Self-employed fallback
        if current_company == "Not Provided":
            if re.search(r'\b(freelanc|self.?employ|self employ|independent consultant|own business|proprietor)\b', post_exp_text, re.IGNORECASE):
                current_company = "Freelance / Self-Employed"

        for o in other_potential:
            if o.lower() != current_company.lower() and o.lower() not in [p.lower() for p in previous_list]:
                previous_list.append(o)

        data["Current Company"] = current_company
        if previous_list:
            data["Previous Companies"] = ", ".join(previous_list[:5])


    return data


def process_file(file_path):
    file_name = os.path.basename(file_path)
    try:
        if file_path.lower().endswith(".docx"):
            extracted_text, hidden_emails = extract_docx_data(file_path)
        else:
            extracted_text, hidden_emails = extract_pdf_data(file_path)
            
        res = parse_text(extracted_text, hidden_emails, file_name)

        
        # Last check - if name is still empty, use sanitized file name
        if res.get('Candidate Name') == "Not Provided":
            res['Candidate Name'] = sanitize_filename(file_name)
            
        return res
    except Exception as e:
        print(f"Processing Error ({file_name}): {e}")
        # Return empty data with just the file name
        res = {k: "Not Provided" for k in FINAL_HEADERS}
        res['File Name'] = file_name
        res['Candidate Name'] = f"Error: {str(e)[:30]}"
        return res


def process_all(files):
    if not files: return []
    # Swap to ThreadPool for pure speed and lightweight execution
    workers = min(16, multiprocessing.cpu_count() * 4)
    with ThreadPoolExecutor(max_workers=workers) as executor:
        results = list(executor.map(process_file, files))
    return results


# ==========================================
# WEB UI & ROUTES
# ==========================================

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        if "zipfile" not in request.files:
            return "No file uploaded", 400
            
        file = request.files["zipfile"]
        if file.filename == "":
            return "No file selected", 400

        zip_path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(zip_path)

        try:
            folder = unzip_file(zip_path)
            if not folder:
                return "Invalid ZIP file format.", 400

            files = get_files(folder)
            print(f"📦 Found {len(files)} resumes in ZIP. Starting extraction...")
            results = process_all(files)
            print(f"✅ Finished processing {len(results)} resumes. Generating Excel...")

            if results:
                df = pd.DataFrame(results)[FINAL_HEADERS]
            else:
                df = pd.DataFrame(columns=FINAL_HEADERS)

            output_file = os.path.join(OUTPUT_FOLDER, "processed_resumes.xlsx")
            df.to_excel(output_file, index=False)

            # Apply neat styling to match requested layout
            from openpyxl import load_workbook
            from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
            from openpyxl.utils import get_column_letter

            wb = load_workbook(output_file)
            ws = wb.active

            header_fill = PatternFill(start_color="1B5586", end_color="1B5586", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True)
            center_align = Alignment(horizontal="center", vertical="center")
            
            thin_border = Border(
                left=Side(style='thin', color='000000'),
                right=Side(style='thin', color='000000'),
                top=Side(style='thin', color='000000'),
                bottom=Side(style='thin', color='000000')
            )

            # Format Headers
            for col_num, cell in enumerate(ws[1], 1):
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = center_align
                cell.border = thin_border
                
                # Auto-adjust width (roughly)
                col_letter = get_column_letter(col_num)
                ws.column_dimensions[col_letter].width = 26

            # Format Data rows (borders & alternating colors)
            stripe_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
            for row_idx, row in enumerate(ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=ws.max_column), 2):
                for cell in row:
                    cell.border = thin_border
                    cell.alignment = Alignment(vertical="center", wrap_text=True)
                    if row_idx % 2 == 0:
                        cell.fill = stripe_fill

            wb.save(output_file)

            shutil.rmtree(folder, ignore_errors=True)
            if os.path.exists(zip_path):
                os.remove(zip_path)

            return send_file(output_file, as_attachment=True)
            
        except Exception as e:
            return f"An error occurred during processing: {str(e)}", 500

    return render_template("index.html")


if __name__ == "__main__":
    # Light start
    print("Resume Extraction Engine Started 🚀")
    app.run(debug=True, port=5000)




